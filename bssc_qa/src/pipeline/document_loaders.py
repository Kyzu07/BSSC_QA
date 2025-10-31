"""Document loaders for multiple formats."""
from pathlib import Path
from typing import Dict, Any
import re

def load_txt(file_path: Path) -> Dict[str, Any]:
    """Load text file."""
    with open(file_path, 'r', encoding='utf-8') as f:
        content = f.read()
    
    return {
        'content': content,
        'metadata': {
            'source': str(file_path),
            'filename': file_path.name,
            'format': 'txt'
        }
    }

def load_pdf(file_path: Path) -> Dict[str, Any]:
    """Load PDF file."""
    from pypdf import PdfReader
    
    reader = PdfReader(file_path)
    content = ''
    
    for page in reader.pages:
        content += page.extract_text() + '\n\n'
    
    return {
        'content': content,
        'metadata': {
            'source': str(file_path),
            'filename': file_path.name,
            'format': 'pdf',
            'pages': len(reader.pages)
        }
    }

def load_html(file_path: Path) -> Dict[str, Any]:
    """Load HTML file."""
    from bs4 import BeautifulSoup
    
    with open(file_path, 'r', encoding='utf-8') as f:
        html_content = f.read()
    
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.decompose()
    
    content = soup.get_text()
    
    return {
        'content': content,
        'metadata': {
            'source': str(file_path),
            'filename': file_path.name,
            'format': 'html'
        }
    }

def load_docx(file_path: Path) -> Dict[str, Any]:
    """Load DOCX file."""
    from docx import Document
    
    doc = Document(file_path)
    content = '\n\n'.join([para.text for para in doc.paragraphs if para.text.strip()])
    
    return {
        'content': content,
        'metadata': {
            'source': str(file_path),
            'filename': file_path.name,
            'format': 'docx',
            'paragraphs': len(doc.paragraphs)
        }
    }

def load_url(url: str) -> Dict[str, Any]:
    """Load content from URL."""
    import requests
    from bs4 import BeautifulSoup
    
    response = requests.get(url, timeout=10)
    response.raise_for_status()
    
    soup = BeautifulSoup(response.content, 'html.parser')
    
    # Remove script and style elements
    for script in soup(["script", "style"]):
        script.decompose()
    
    content = soup.get_text()
    
    return {
        'content': content,
        'metadata': {
            'source': url,
            'format': 'url'
        }
    }

def load_document(source: str) -> Dict[str, Any]:
    """Load document from file path or URL."""
    # Check if URL
    if source.startswith('http://') or source.startswith('https://'):
        return load_url(source)
    
    # Load from file
    file_path = Path(source)
    
    if not file_path.exists():
        raise FileNotFoundError(f"File not found: {source}")
    
    suffix = file_path.suffix.lower()
    
    loaders = {
        '.txt': load_txt,
        '.pdf': load_pdf,
        '.html': load_html,
        '.htm': load_html,
        '.docx': load_docx
    }
    
    loader = loaders.get(suffix)
    if loader is None:
        raise ValueError(f"Unsupported file format: {suffix}")
    
    return loader(file_path)